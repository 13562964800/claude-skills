#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
自动插入所有插图到Word文档
将SVG转换为PNG并插入到对应位置
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re
import subprocess
import os

def svg_to_png(svg_path, png_path, width=800):
    """将SVG转换为PNG"""
    try:
        # 使用cairosvg转换（如果安装了）
        import cairosvg
        cairosvg.svg2png(url=str(svg_path), write_to=str(png_path), output_width=width)
        return True
    except ImportError:
        print(f"  ⚠️  cairosvg未安装，尝试使用其他方法...")

    try:
        # 尝试使用Inkscape
        cmd = ['inkscape', str(svg_path), '--export-png', str(png_path), f'--export-width={width}']
        subprocess.run(cmd, check=True, capture_output=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        print(f"  ⚠️  Inkscape未安装")

    return False

def find_illustration_files():
    """查找所有插图文件"""
    illustration_dir = Path("C:/Users/Administrator/插图")
    svg_files = list(illustration_dir.glob("*.svg"))

    # 创建插图映射
    illustration_map = {}
    for svg_file in svg_files:
        name = svg_file.stem
        illustration_map[name] = svg_file

    return illustration_map

def insert_illustrations(doc_path, output_path):
    """插入所有插图"""
    print("📖 加载Word文档...")
    doc = Document(str(doc_path))

    print("🔍 查找插图文件...")
    illustration_map = find_illustration_files()
    print(f"  ✓ 找到 {len(illustration_map)} 个插图文件")

    # 创建PNG临时目录
    png_dir = Path("C:/Users/Administrator/插图/png")
    png_dir.mkdir(exist_ok=True)

    print("\n🎨 转换SVG为PNG...")
    png_map = {}
    for name, svg_path in illustration_map.items():
        png_path = png_dir / f"{name}.png"
        if not png_path.exists():
            print(f"  转换: {name}.svg → {name}.png")
            if svg_to_png(svg_path, png_path):
                png_map[name] = png_path
            else:
                # 如果转换失败，直接使用SVG
                png_map[name] = svg_path
        else:
            png_map[name] = png_path

    print(f"\n🖼️  插入插图到文档...")
    inserted_count = 0
    placeholder_pattern = re.compile(r'\[插图占位符 #\d+\]')

    # 遍历所有段落
    paragraphs_to_process = []
    for i, para in enumerate(doc.paragraphs):
        if placeholder_pattern.search(para.text):
            paragraphs_to_process.append((i, para))

    print(f"  找到 {len(paragraphs_to_process)} 个插图占位符")

    # 处理每个占位符
    for idx, (para_idx, para) in enumerate(paragraphs_to_process):
        # 尝试从上下文找到插图名称
        # 查找前后几个段落中的【插图X-X】标记
        context_start = max(0, para_idx - 3)
        context_end = min(len(doc.paragraphs), para_idx + 3)

        illustration_name = None
        for j in range(context_start, context_end):
            text = doc.paragraphs[j].text
            # 匹配【插图X-X：名称】或类似格式
            match = re.search(r'【插图[^】]+】', text)
            if match:
                # 提取插图描述，尝试匹配文件名
                desc = match.group(0)
                # 尝试在插图映射中找到匹配的文件
                for name in png_map.keys():
                    if name in desc or any(part in name for part in desc.split('：')[-1].split('-')):
                        illustration_name = name
                        break
                if illustration_name:
                    break

        # 如果找到了对应的插图文件
        if illustration_name and illustration_name in png_map:
            try:
                # 清空占位符段落
                para.clear()

                # 插入图片
                run = para.add_run()
                run.add_picture(str(png_map[illustration_name]), width=Inches(5.5))

                # 居中对齐
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                inserted_count += 1
                print(f"  ✓ [{inserted_count}/{len(paragraphs_to_process)}] 插入: {illustration_name}")
            except Exception as e:
                print(f"  ✗ 插入失败: {illustration_name} - {e}")
        else:
            # 如果没找到对应文件，使用通用插图
            # 尝试使用序号匹配
            placeholder_num = idx + 1
            if placeholder_num <= len(png_map):
                try:
                    available_illustrations = list(png_map.values())
                    para.clear()
                    run = para.add_run()
                    run.add_picture(str(available_illustrations[placeholder_num % len(available_illustrations)]), width=Inches(5.5))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    inserted_count += 1
                    print(f"  ✓ [{inserted_count}/{len(paragraphs_to_process)}] 插入通用插图 #{placeholder_num}")
                except Exception as e:
                    print(f"  ✗ 插入失败 #{placeholder_num} - {e}")

    print(f"\n💾 保存文档...")
    doc.save(str(output_path))
    print(f"  ✓ 已保存: {output_path}")

    return inserted_count

def main():
    """主函数"""
    print("=" * 70)
    print("🎨 自动插入所有插图到Word文档")
    print("=" * 70)

    input_file = Path("C:/Users/Administrator/AI科普画册/AI前沿知识儿童科普画册-完整版-带插图占位符.docx")
    output_file = Path("C:/Users/Administrator/AI科普画册/AI前沿知识儿童科普画册-完整版-已插入插图.docx")

    if not input_file.exists():
        print(f"❌ 找不到输入文件: {input_file}")
        return

    print(f"\n📂 输入文件: {input_file}")
    print(f"📂 输出文件: {output_file}\n")

    try:
        inserted_count = insert_illustrations(input_file, output_file)

        print("\n" + "=" * 70)
        print(f"✅ 完成！共插入 {inserted_count} 个插图")
        print(f"📁 输出文件: {output_file}")
        print("=" * 70)

    except Exception as e:
        print(f"\n❌ 错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
