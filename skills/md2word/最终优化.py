#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI科普画册最终整合与优化脚本
完成所有剩余的整合工作
"""

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

def add_cover_page(doc, cover_svg_path):
    """添加封面页"""
    print("📕 添加封面...")
    # 在文档开头插入封面
    # 注意：SVG需要转换为图片格式才能插入Word
    print("  ⚠️  封面需要手动插入（SVG转PNG后）")

def fix_font_consistency(doc):
    """修正字体一致性"""
    print("🎨 修正字体一致性...")

    # 定义标准颜色
    COLORS = {
        "heading1": RGBColor(41, 128, 185),    # 蓝色
        "heading2": RGBColor(142, 68, 173),    # 紫色
        "heading3": RGBColor(39, 174, 96),     # 绿色
        "normal": RGBColor(0, 0, 0),           # 黑色
    }

    fixed_count = 0

    for para in doc.paragraphs:
        style_name = para.style.name.lower()

        # 确定目标颜色
        if "heading 1" in style_name or "标题 1" in style_name:
            target_color = COLORS["heading1"]
        elif "heading 2" in style_name or "标题 2" in style_name:
            target_color = COLORS["heading2"]
        elif "heading 3" in style_name or "标题 3" in style_name:
            target_color = COLORS["heading3"]
        else:
            target_color = COLORS["normal"]

        # 修正所有文本块的颜色
        for run in para.runs:
            if run.font.color.rgb != target_color:
                run.font.color.rgb = target_color
                fixed_count += 1

    print(f"  ✓ 修正了 {fixed_count} 处字体颜色")

def insert_illustration_placeholders(doc):
    """为插图标记添加占位符"""
    print("🖼️  处理插图标记...")

    illustration_count = 0

    for para in doc.paragraphs:
        # 查找【插图X-X】标记
        if "【插图" in para.text or "[插图" in para.text:
            illustration_count += 1

            # 添加占位符段落
            new_para = para.insert_paragraph_before()
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加占位文本
            run = new_para.add_run(f"\n[插图占位符 #{illustration_count}]\n")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.italic = True

    print(f"  ✓ 找到 {illustration_count} 个插图位置")

def add_page_numbers(doc):
    """添加页码"""
    print("📄 配置页码...")
    print("  ⚠️  页码需要在Word中手动添加")

def update_toc(doc):
    """更新目录"""
    print("📑 更新目录...")
    print("  ⚠️  目录需要在Word中手动更新（右键→更新域）")

def apply_comic_style(doc):
    """应用漫画书风格"""
    print("🎨 应用漫画书风格...")

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Comic Sans MS'
    font.size = Pt(14)

    # 设置行距
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 2.0

    print("  ✓ 已应用Comic Sans MS字体和2倍行距")

def main():
    """主函数"""
    print("=" * 70)
    print("🎨 AI科普画册最终整合与优化")
    print("=" * 70)

    # 文件路径
    input_file = Path("C:/Users/Administrator/AI科普画册/AI前沿知识儿童科普画册-最终整合版.docx")
    output_file = Path("C:/Users/Administrator/AI科普画册/AI前沿知识儿童科普画册-完整版-带插图占位符.docx")
    cover_svg = Path("C:/Users/Administrator/插图/封面-精美版.svg")

    if not input_file.exists():
        print(f"❌ 找不到输入文件: {input_file}")
        return

    print(f"\n📂 输入文件: {input_file}")
    print(f"📂 输出文件: {output_file}\n")

    # 加载文档
    print("📖 加载Word文档...")
    doc = Document(str(input_file))
    print(f"  ✓ 文档已加载，共 {len(doc.paragraphs)} 个段落\n")

    # 执行各项优化
    apply_comic_style(doc)
    fix_font_consistency(doc)
    insert_illustration_placeholders(doc)
    add_page_numbers(doc)
    update_toc(doc)

    # 保存文档
    print(f"\n💾 保存文档...")
    doc.save(str(output_file))
    print(f"  ✓ 已保存到: {output_file}")

    # 生成后续步骤说明
    print("\n" + "=" * 70)
    print("✅ 自动化处理完成！")
    print("\n📋 后续手动步骤：")
    print("\n1. 插入插图：")
    print("   - 打开Word文档")
    print("   - 在每个[插图占位符]位置插入对应的SVG图片")
    print("   - 调整图片大小和位置")
    print("\n2. 添加封面封底：")
    print("   - 将封面-精美版.svg转换为PNG")
    print("   - 插入到文档第一页")
    print("   - 将封底-精美版.svg转换为PNG")
    print("   - 插入到文档最后一页")
    print("\n3. 重新计算页码：")
    print("   - 插入→页码→页面底端→居中")
    print("   - 设置封面不显示页码")
    print("   - 目录页使用罗马数字（i, ii, iii...）")
    print("   - 正文页使用阿拉伯数字（1, 2, 3...）")
    print("\n4. 更新目录：")
    print("   - 右键点击目录")
    print("   - 选择\"更新域\"")
    print("   - 选择\"更新整个目录\"")
    print("\n5. 最终检查：")
    print("   - 检查所有插图是否正确显示")
    print("   - 检查页码是否连续")
    print("   - 检查目录页码是否对应")
    print("   - 检查字体颜色是否一致")
    print("   - 检查排版是否美观")
    print("\n" + "=" * 70)

if __name__ == "__main__":
    main()
