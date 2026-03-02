#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将AI科普画册Word文档优化为漫画书风格
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys

def set_cell_border(cell, **kwargs):
    """设置表格单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), str(edge_data.get('sz', 12)))
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def create_comic_style_doc(input_path, output_path):
    """创建漫画书风格的Word文档"""

    print(f"正在读取文档: {input_path}")
    doc = Document(input_path)

    # 创建新文档
    new_doc = Document()

    # 设置页面边距（更大的留白）
    sections = new_doc.sections
    for section in sections:
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(1.2)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 定义漫画书风格样式
    styles = new_doc.styles

    # 标题1样式（章节标题）- 蓝色
    if 'Comic Heading 1' not in [s.name for s in styles]:
        heading1_style = styles.add_style('Comic Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.font.name = 'Comic Sans MS'
        heading1_style.font.size = Pt(28)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色
        heading1_style.paragraph_format.space_before = Pt(24)
        heading1_style.paragraph_format.space_after = Pt(18)
        heading1_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 标题2样式（小节标题）- 粉色
    if 'Comic Heading 2' not in [s.name for s in styles]:
        heading2_style = styles.add_style('Comic Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.font.name = 'Comic Sans MS'
        heading2_style.font.size = Pt(22)
        heading2_style.font.bold = True
        heading2_style.font.color.rgb = RGBColor(255, 51, 153)  # 粉色
        heading2_style.paragraph_format.space_before = Pt(18)
        heading2_style.paragraph_format.space_after = Pt(12)
        heading2_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # 标题3样式（子标题）- 绿色
    if 'Comic Heading 3' not in [s.name for s in styles]:
        heading3_style = styles.add_style('Comic Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        heading3_style.font.name = 'Comic Sans MS'
        heading3_style.font.size = Pt(18)
        heading3_style.font.bold = True
        heading3_style.font.color.rgb = RGBColor(0, 153, 76)  # 绿色
        heading3_style.paragraph_format.space_before = Pt(12)
        heading3_style.paragraph_format.space_after = Pt(8)
        heading3_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # 正文样式
    if 'Comic Body' not in [s.name for s in styles]:
        body_style = styles.add_style('Comic Body', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Comic Sans MS'
        body_style.font.size = Pt(14)
        body_style.paragraph_format.line_spacing = 2.0  # 2倍行距
        body_style.paragraph_format.space_after = Pt(12)
        body_style.paragraph_format.first_line_indent = Inches(0.3)

    # 插图说明样式
    if 'Comic Caption' not in [s.name for s in styles]:
        caption_style = styles.add_style('Comic Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = 'Comic Sans MS'
        caption_style.font.size = Pt(12)
        caption_style.font.italic = True
        caption_style.font.color.rgb = RGBColor(102, 102, 102)  # 灰色
        caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_style.paragraph_format.space_after = Pt(18)

    print("正在处理文档内容...")

    # 处理原文档内容
    page_count = 1
    current_section = None

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        # 检测标题级别
        if text.startswith('# '):
            # 一级标题 - 封面或主标题
            new_para = new_doc.add_paragraph(text[2:])
            new_para.style = 'Comic Heading 1'
            new_doc.add_page_break()
            page_count += 1

        elif text.startswith('## '):
            # 二级标题 - 章节标题
            if current_section:
                new_doc.add_page_break()  # 每章从新页开始
                page_count += 1
            new_para = new_doc.add_paragraph(text[3:])
            new_para.style = 'Comic Heading 1'
            current_section = text[3:]

        elif text.startswith('### '):
            # 三级标题 - 小节标题
            new_para = new_doc.add_paragraph(text[4:])
            new_para.style = 'Comic Heading 2'

        elif text.startswith('#### '):
            # 四级标题 - 子标题
            new_para = new_doc.add_paragraph(text[5:])
            new_para.style = 'Comic Heading 3'

        elif '【插图' in text or '【封面插图】' in text:
            # 插图标记 - 预留空间
            new_doc.add_paragraph()  # 空行

            # 添加插图占位框
            table = new_doc.add_table(rows=1, cols=1)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell = table.rows[0].cells[0]
            cell.width = Inches(5.0)
            cell_para = cell.paragraphs[0]
            cell_para.text = text
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 设置边框
            set_cell_border(
                cell,
                top={"sz": 12, "color": "CCCCCC"},
                left={"sz": 12, "color": "CCCCCC"},
                bottom={"sz": 12, "color": "CCCCCC"},
                right={"sz": 12, "color": "CCCCCC"}
            )

            # 添加插图说明
            caption = new_doc.add_paragraph(text)
            caption.style = 'Comic Caption'
            new_doc.add_paragraph()  # 空行

        elif text.startswith('**') and text.endswith('**'):
            # 粗体文本
            new_para = new_doc.add_paragraph()
            run = new_para.add_run(text[2:-2])
            run.bold = True
            run.font.name = 'Comic Sans MS'
            run.font.size = Pt(14)
            new_para.paragraph_format.line_spacing = 2.0

        elif text.startswith('- ') or text.startswith('* '):
            # 列表项
            new_para = new_doc.add_paragraph(text[2:], style='List Bullet')
            new_para.style.font.name = 'Comic Sans MS'
            new_para.style.font.size = Pt(14)
            new_para.paragraph_format.line_spacing = 2.0
            new_para.paragraph_format.left_indent = Inches(0.5)

        elif text.startswith('---'):
            # 分隔线 - 添加分隔页
            new_doc.add_page_break()
            page_count += 1

        else:
            # 普通段落
            new_para = new_doc.add_paragraph(text)
            new_para.style = 'Comic Body'

    # 添加页眉页脚
    section = new_doc.sections[0]

    # 页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "AI前沿知识儿童科普画册"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.style.font.name = 'Comic Sans MS'
    header_para.style.font.size = Pt(10)
    header_para.style.font.color.rgb = RGBColor(128, 128, 128)

    # 页脚（页码）
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.style.font.name = 'Comic Sans MS'
    footer_para.style.font.size = Pt(10)

    print(f"正在保存文档: {output_path}")
    new_doc.save(output_path)

    return page_count

def main():
    # 文件路径
    home = os.path.expanduser('~')
    input_file = os.path.join(home, 'AI科普画册', 'AI前沿知识儿童科普画册-漫画版.docx')
    output_file = os.path.join(home, 'AI科普画册', 'AI前沿知识儿童科普画册-漫画书.docx')

    # 检查输入文件
    if not os.path.exists(input_file):
        print(f"输入文件不存在: {input_file}")
        print("尝试使用完整版.md生成...")
        md_file = os.path.join(home, 'AI科普画册', 'AI前沿知识儿童科普画册-完整版.md')
        if not os.path.exists(md_file):
            print(f"Markdown文件也不存在: {md_file}")
            sys.exit(1)
        # 这里可以添加从MD生成DOCX的逻辑
        print("请先从MD生成DOCX文件")
        sys.exit(1)

    print("=" * 60)
    print("AI科普画册 - 漫画书风格优化工具")
    print("=" * 60)

    try:
        page_count = create_comic_style_doc(input_file, output_file)

        print("\n" + "=" * 60)
        print("✓ 文档优化完成！")
        print("=" * 60)
        print(f"\n输出文件: {output_file}")
        print(f"预估页数: {page_count} 页")
        print("\n格式优化说明：")
        print("  ✓ 使用Comic Sans MS活泼字体")
        print("  ✓ 2倍行距，留白充足")
        print("  ✓ 章节标题彩色（蓝色/粉色/绿色）")
        print("  ✓ 每章从新页开始")
        print("  ✓ 插图位置预留空间")
        print("  ✓ 添加页眉页脚")
        print("  ✓ 居中对齐的插图说明")
        print("\n" + "=" * 60)

    except Exception as e:
        print(f"\n错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()
